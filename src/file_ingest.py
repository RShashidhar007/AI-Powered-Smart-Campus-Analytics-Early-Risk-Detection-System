"""
file_ingest.py - Extract student data tables from uploaded PDFs, CSVs, and Excel files.
Uses pdfplumber for table extraction, and pandas for CSV/Excel reading. Handles missing columns gracefully.
"""
import pandas as pd
import pdfplumber
import io
import re
from docx import Document


# Expected columns in our dataset
EXPECTED_COLUMNS = [
    'usn', 'name', 'department', 'semester', 'attendance', 'internal_marks',
    'assignment_score', 'quiz_score', 'lab_marks',
    'semester_marks', 'study_hours',
]

# Common aliases -> canonical name mapping
COLUMN_ALIASES = {
    # USN
    'usn': 'usn', 'roll_no': 'usn', 'roll': 'usn', 'rollno': 'usn',
    'student_id': 'usn', 'id': 'usn', 'reg_no': 'usn', 'registration': 'usn',
    'enrollment': 'usn', 'enrollment_no': 'usn',
    # Name
    'name': 'name', 'student_name': 'name', 'student': 'name', 'full_name': 'name',
    # Attendance
    'attendance': 'attendance', 'attend': 'attendance', 'attendance_%': 'attendance',
    'attendance_percent': 'attendance', 'att': 'attendance', 'att%': 'attendance',
    # Internal marks
    'internal_marks': 'internal_marks', 'internal': 'internal_marks',
    'internals': 'internal_marks', 'cie': 'internal_marks',
    'cie_marks': 'internal_marks', 'ia': 'internal_marks', 'ia_marks': 'internal_marks',
    # Assignment
    'assignment_score': 'assignment_score', 'assignment': 'assignment_score',
    'assignments': 'assignment_score', 'asgn': 'assignment_score',
    # Quiz
    'quiz_score': 'quiz_score', 'quiz': 'quiz_score', 'quiz_marks': 'quiz_score',
    'quizzes': 'quiz_score',
    # Lab
    'lab_marks': 'lab_marks', 'lab': 'lab_marks', 'practical': 'lab_marks',
    'lab_score': 'lab_marks', 'practical_marks': 'lab_marks',
    # Semester marks
    'semester_marks': 'semester_marks', 'sem_marks': 'semester_marks',
    'semester': 'semester_marks', 'total_marks': 'semester_marks',
    'exam_marks': 'semester_marks', 'see': 'semester_marks',
    'see_marks': 'semester_marks', 'final_marks': 'semester_marks',
    'marks': 'semester_marks', 'total': 'semester_marks',
    # Study hours
    'study_hours': 'study_hours', 'study': 'study_hours',
    'hours': 'study_hours', 'study_hrs': 'study_hours',
    # Department
    'department': 'department', 'dept': 'department', 'branch': 'department',
    'dept_name': 'department', 'department_name': 'department',
    # Semester
    'semester': 'semester', 'sem': 'semester', 'sem_no': 'semester',
    'semester_no': 'semester', 'term': 'semester',
}


def _normalise_col(col: str) -> str:
    """Normalise a column name to lowercase, underscored, stripped."""
    col = str(col).strip().lower()
    col = re.sub(r'[^a-z0-9%]', '_', col)
    col = re.sub(r'_+', '_', col).strip('_')
    return col


def _map_columns(raw_cols: list[str]) -> dict[str, str]:
    """Map raw column names to canonical column names."""
    mapping = {}
    for raw in raw_cols:
        norm = _normalise_col(raw)
        if norm in COLUMN_ALIASES:
            canonical = COLUMN_ALIASES[norm]
            if canonical not in mapping.values():  # avoid duplicates
                mapping[raw] = canonical
    return mapping


def extract_tables_from_pdf(file_bytes: bytes) -> tuple[pd.DataFrame | None, str]:
    """
    Extract student data from a PDF.

    Returns:
        (DataFrame or None, status_message)
    """
    try:
        pdf = pdfplumber.open(io.BytesIO(file_bytes))
    except Exception as e:
        return None, f" Could not open PDF: {e}", False

    all_rows = []
    header = None

    for page_num, page in enumerate(pdf.pages, 1):
        tables = page.extract_tables()
        if not tables:
            continue

        for table in tables:
            if not table or len(table) < 2:
                continue

            # First row of first table is likely the header
            if header is None:
                header = [str(c).strip() if c else f"col_{i}" for i, c in enumerate(table[0])]
                data_rows = table[1:]
            else:
                data_rows = table
                # Skip rows that look like repeated headers
                data_rows = [r for r in data_rows
                             if r and str(r[0]).strip().lower() != str(header[0]).strip().lower()]

            for row in data_rows:
                if row and any(cell and str(cell).strip() for cell in row):
                    # Pad or truncate row to match header length
                    padded = (list(row) + [None] * len(header))[:len(header)]
                    all_rows.append(padded)

    pdf.close()

    if not header or not all_rows:
        return None, " No tables found in the PDF. Make sure the PDF contains tabular student data.", False

    # Build DataFrame
    raw_df = pd.DataFrame(all_rows, columns=header)
    return _clean_and_map_dataframe(raw_df, "PDF")


def extract_tables_from_docx(file_bytes: bytes) -> tuple[pd.DataFrame | None, str, bool]:
    """
    Extract student data from a Word document (.docx).
    Reads all tables found in the document and combines them.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        return None, f" Could not open Word document: {e}", False

    if not doc.tables:
        return None, " No tables found in the Word document. Make sure it contains tabular student data.", False

    all_rows = []
    header = None

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            if header is None:
                header = cells
            else:
                # Skip repeated header rows
                if cells == header:
                    continue
                all_rows.append(cells)

    if not header or not all_rows:
        return None, " Could not extract data from Word document tables.", False

    raw_df = pd.DataFrame(all_rows, columns=header)
    return _clean_and_map_dataframe(raw_df, "Word Document")


def process_uploaded_file(file_bytes: bytes, file_name: str) -> tuple[pd.DataFrame | None, str]:
    """
    Extract student data from an uploaded file (PDF, CSV, XLSX).
    """
    ext = file_name.lower().split('.')[-1]
    
    if ext == 'pdf':
        return extract_tables_from_pdf(file_bytes)
    elif ext == 'csv':
        try:
            raw_df = pd.read_csv(io.BytesIO(file_bytes))
            return _clean_and_map_dataframe(raw_df, "CSV")
        except Exception as e:
            return None, f" Could not read CSV: {e}", False
    elif ext in ['xlsx', 'xls']:
        try:
            raw_df = pd.read_excel(io.BytesIO(file_bytes))
            return _clean_and_map_dataframe(raw_df, "Excel")
        except Exception as e:
            return None, f" Could not read Excel: {e}", False
    elif ext in ['docx', 'doc']:
        return extract_tables_from_docx(file_bytes)
    else:
        return None, f" Unsupported file type: {ext}", False


def _clean_and_map_dataframe(raw_df: pd.DataFrame, source_type: str) -> tuple[pd.DataFrame | None, str]:
    header = list(raw_df.columns)
    
    # Map columns
    col_mapping = _map_columns(header)
    if not col_mapping:
        return None, (
            f" Could not map any columns in the {source_type}. Found: **{', '.join(str(c) for c in header)}**\n\n"
            f"Expected columns like: usn, name, attendance, internal_marks, "
            f"assignment_score, quiz_score, lab_marks, semester_marks, study_hours"
        ), False

    df = raw_df.rename(columns=col_mapping)

    # Keep only mapped columns
    mapped_cols = [c for c in EXPECTED_COLUMNS if c in df.columns]
    df = df[mapped_cols].copy()

    # Convert numeric columns
    numeric_cols = [c for c in mapped_cols if c not in ('usn', 'name')]
    for col in numeric_cols:
        df[col] = pd.to_numeric(df[col], errors='coerce')

    # Drop rows that are entirely NaN (bad parses)
    df = df.dropna(how='all').reset_index(drop=True)

    # Fill missing numeric columns with sensible defaults
    missing = [c for c in EXPECTED_COLUMNS if c not in mapped_cols]
    defaults = {
        'usn': 'UNKNOWN', 'name': 'Unknown Student',
        'semester': 1,
        'attendance': 75.0, 'internal_marks': 30.0,
        'assignment_score': 30.0, 'quiz_score': 30.0,
        'lab_marks': 30.0, 'semester_marks': 140.0, 'study_hours': 3.0,
    }
    for col in missing:
        df[col] = defaults.get(col, 0)

    # Flag if department is missing or has empty values — UI will ask the user to pick one
    if 'department' in missing:
        dept_missing = True
    elif 'department' in df.columns:
        # Column exists but values are empty/null
        empty_dept = df['department'].isna() | (df['department'].astype(str).str.strip() == '') | (df['department'].astype(str).str.strip() == '0')
        if empty_dept.any():
            dept_missing = True
            df.loc[empty_dept, 'department'] = ''
        else:
            dept_missing = False
    else:
        dept_missing = False

    # Auto-generate USN if missing
    if 'usn' in missing:
        df['usn'] = [f"NEW{i+1:04d}" for i in range(len(df))]

    # Ensure correct column order
    df = df[EXPECTED_COLUMNS]

    # Build status message
    found_str = ", ".join(f"**{c}**" for c in mapped_cols)
    status = f" Extracted **{len(df)} students** from {source_type}.\n\n Mapped columns: {found_str}"
    if missing:
        missing_str = ", ".join(f"`{c}`" for c in missing)
        status += f"\n\n Missing columns (filled with defaults): {missing_str}"

    return df, status, dept_missing
